"""
Generador d'informes mensuals — Adtende Analytics (format .docx)
"""

import io
import calendar
from datetime import date
from pathlib import Path

import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from api_client import AdtendeClient


# ── Paleta ───────────────────────────────────────────────────────────────────
C_NAVY  = RGBColor(0x1B, 0x3A, 0x6B)
C_TEAL  = RGBColor(0x00, 0xB4, 0xA6)
C_AMBER = RGBColor(0xF4, 0xA6, 0x20)
C_RED   = RGBColor(0xE0, 0x5A, 0x5A)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_LIGHT = RGBColor(0xF4, 0xF7, 0xFB)
C_MID   = RGBColor(0xDD, 0xE6, 0xF0)
C_TEXT  = RGBColor(0x2C, 0x2C, 0x2C)
C_MUTED = RGBColor(0x7A, 0x8F, 0xA6)

# Hex strings for matplotlib
H_NAVY  = "#1B3A6B"
H_TEAL  = "#00B4A6"
H_AMBER = "#F4A620"
H_RED   = "#E05A5A"
H_LIGHT = "#F4F7FB"
H_MID   = "#DDE6F0"
H_TEXT  = "#2C2C2C"
H_MUTED = "#7A8FA6"

CHART_COLORS = [H_NAVY, H_TEAL, H_AMBER, H_RED, "#7B68EE", "#52C784", "#F47560", "#97BBD5"]

MESOS = {
    1:"Gener", 2:"Febrer", 3:"Març", 4:"Abril",
    5:"Maig", 6:"Juny", 7:"Juliol", 8:"Agost",
    9:"Setembre", 10:"Octubre", 11:"Novembre", 12:"Desembre",
}
DIES = ["Dilluns","Dimarts","Dimecres","Dijous","Divendres","Dissabte","Diumenge"]


# ── Helpers numèrics ──────────────────────────────────────────────────────────

def _fmt(n):
    if n is None: return "—"
    if isinstance(n, float):
        return f"{n:,.1f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return f"{int(n):,}".replace(",", ".")

def _pct(a, b): return round(a / b * 100, 1) if b else 0.0

def _fmt_time(minutes):
    if minutes is None: return "—"
    total_s = round(minutes * 60)
    m, s = divmod(total_s, 60)
    return f"{m}m {s:02d}s"


# ── Helpers docx ─────────────────────────────────────────────────────────────

def _hex(val) -> str:
    """Retorna hex pur (sense #) per a atributs XML de docx.
    Accepta RGBColor (tuple r,g,b) o string '#RRGGBB'/'RRGGBB'.
    """
    if isinstance(val, str):
        return val.replace("#", "").upper()
    # RGBColor és una tuple (r, g, b)
    try:
        r, g, b = val[0], val[1], val[2]
        return f"{r:02X}{g:02X}{b:02X}"
    except Exception:
        return "000000"

def _set_cell_bg(cell, hex_color: str):
    """Posa color de fons a una cel·la de taula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Afegeix vores a una cel·la."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), str(val.get("sz", 4)))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", "000000"))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def _run(para, text, bold=False, italic=False, size=None, color=None, font="Calibri"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:   run.font.size = Pt(size)
    if color:  run.font.color.rgb = color
    run.font.name = font
    return run

def _para(doc_or_cell, text="", bold=False, size=None, color=None,
          align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    if hasattr(doc_or_cell, "add_paragraph"):
        p = doc_or_cell.add_paragraph()
    else:
        p = doc_or_cell.paragraphs[0] if doc_or_cell.paragraphs else doc_or_cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        _run(p, text, bold=bold, size=size, color=color)
    return p

def _section_title(doc, num, title):
    """Capçalera de secció estilitzada."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    # Número teal
    r1 = p.add_run(f"{num}  ")
    r1.font.bold = True
    r1.font.size = Pt(13)
    r1.font.color.rgb = C_TEAL
    r1.font.name = "Calibri"
    # Títol navy
    r2 = p.add_run(title)
    r2.font.bold = True
    r2.font.size = Pt(13)
    r2.font.color.rgb = C_NAVY
    r2.font.name = "Calibri"
    # Línia horitzontal sota
    _add_hrule(doc, color="00B4A6")
    return p

def _add_hrule(doc, color="DDE6F0", size=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def _body_para(doc, text, bold_parts=None, space_after=6):
    """Paràgraf de cos amb text; bold_parts = [(start, end), ...] (no usat ara)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = C_TEXT
    r.font.name = "Calibri"
    return p


# ── Taules ────────────────────────────────────────────────────────────────────

def _add_table(doc, headers, rows, col_widths_cm=None, total_row=False):
    """
    Afegeix una taula estilitzada al document.
    headers: llista d'strings
    rows: llista de llistes d'strings
    col_widths_cm: llista de floats en cm
    """
    n_cols = len(headers)
    n_rows = len(rows)
    table = doc.add_table(rows=1 + n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Amplades
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    # Capçalera
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        _set_cell_bg(cell, _hex(C_NAVY))
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(h)
        r.font.bold  = True
        r.font.size  = Pt(9)
        r.font.color.rgb = C_WHITE
        r.font.name  = "Calibri"
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Files de dades
    for ri, row_data in enumerate(rows):
        row = table.rows[1 + ri]
        is_total = total_row and ri == n_rows - 1
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            bg = "EBF4FF" if is_total else ("F4F7FB" if ri % 2 == 1 else "FFFFFF")
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            p.alignment = align
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(str(val))
            r.font.bold  = is_total
            r.font.size  = Pt(9)
            r.font.color.rgb = C_NAVY if is_total else C_TEXT
            r.font.name  = "Calibri"
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


# ── KPI cards ─────────────────────────────────────────────────────────────────

def _add_kpi_row(doc, kpis):
    """
    kpis = [(valor_str, label, color_hex), ...]
    Afegeix una fila de targetes KPI com a taula.
    """
    n = len(kpis)
    w = 17.0 / n  # amplada total ~17cm
    table = doc.add_table(rows=2, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (val, lbl, hex_color) in enumerate(kpis):
        # Fila valor
        vc = table.rows[0].cells[i]
        _set_cell_bg(vc, "F4F7FB")
        p = vc.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(val)
        r.font.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = C_NAVY
        r.font.name = "Calibri"
        vc.width = Cm(w)
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_borders(vc, bottom={"val": "single", "sz": 6, "color": hex_color.replace("#", "")})

        # Fila label
        lc = table.rows[1].cells[i]
        _set_cell_bg(lc, "F4F7FB")
        p2 = lc.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after  = Pt(6)
        r2 = p2.add_run(lbl)
        r2.font.size = Pt(8)
        r2.font.color.rgb = C_MUTED
        r2.font.name = "Calibri"
        lc.width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


# ── Gràfiques matplotlib → BytesIO ────────────────────────────────────────────

def _mpl_style():
    plt.rcParams.update({
        "font.family": "sans-serif",
        "axes.spines.top": False,
        "axes.spines.right": False,
        "axes.spines.left": False,
        "axes.grid": True,
        "grid.color": H_MID,
        "grid.linewidth": 0.6,
        "axes.axisbelow": True,
        "xtick.color": H_MUTED,
        "ytick.color": H_MUTED,
        "axes.labelcolor": H_MUTED,
        "figure.facecolor": "white",
    })

def _buf(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf

def _chart_evolucion(df):
    col = "td_managed" if "td_managed" in df.columns else "td_created"
    if col not in df.columns:
        return None
    _mpl_style()
    serie = df.groupby(pd.to_datetime(df[col], errors="coerce").dt.date).size()
    if serie.empty:
        return None
    fig, ax = plt.subplots(figsize=(12, 3.5))
    ax.bar(range(len(serie)), serie.values, color=H_NAVY, alpha=0.85, width=0.7)
    ax.set_xticks(range(len(serie)))
    ax.set_xticklabels([str(d) for d in serie.index], rotation=45, ha="right", fontsize=7)
    ax.set_ylabel("Consultes", fontsize=8)
    ax.set_title("Consultes per dia", fontsize=9, fontweight="bold", color=H_NAVY, pad=6)
    fig.tight_layout(pad=0.6)
    return _buf(fig)

def _chart_heatmap(df):
    if "td_created" not in df.columns or "th_created" not in df.columns:
        return None
    _mpl_style()
    d = df.copy()
    d["td_created"] = pd.to_datetime(d["td_created"], errors="coerce")
    d["dia"] = d["td_created"].dt.dayofweek
    d["hora"] = pd.to_datetime(d["th_created"], format="%H:%M:%S", errors="coerce").dt.hour
    pivot = d.groupby(["dia", "hora"]).size().unstack(fill_value=0)\
             .reindex(index=range(7), columns=range(24), fill_value=0)
    fig, ax = plt.subplots(figsize=(12, 3.6))
    im = ax.imshow(pivot.values, aspect="auto", cmap="Blues", interpolation="nearest", vmin=0)
    ax.set_xticks(range(24))
    ax.set_xticklabels([f"{h:02d}h" for h in range(24)], fontsize=7, rotation=45)
    ax.set_yticks(range(7))
    ax.set_yticklabels(DIES, fontsize=8)
    ax.set_title("Mapa de calor: dia × hora", fontsize=9, fontweight="bold", color=H_NAVY, pad=6)
    # Desactivem completament tota quadrícula
    ax.grid(False)
    ax.tick_params(which="both", length=0)
    for spine in ax.spines.values():
        spine.set_visible(False)
    plt.colorbar(im, ax=ax, shrink=0.8)
    # Anotacions numèriques sense quadrícula
    max_val = pivot.values.max() if pivot.values.max() > 0 else 1
    for row_i in range(7):
        for col_j in range(24):
            val = pivot.values[row_i, col_j]
            if val > 0:
                brightness = val / max_val
                txt_color = "white" if brightness > 0.5 else "#1B3A6B"
                ax.text(col_j, row_i, str(val),
                        ha="center", va="center",
                        fontsize=5.5, color=txt_color, fontweight="bold")
    fig.tight_layout(pad=0.6)
    return _buf(fig)

def _chart_hbars(series, total, color=H_NAVY, top=15):
    s = series.head(top)
    _mpl_style()
    fig, ax = plt.subplots(figsize=(10, max(3, len(s) * 0.55)))
    bars = ax.barh(range(len(s)), s.values, color=color, alpha=0.85, height=0.6)
    ax.set_yticks(range(len(s)))
    ax.set_yticklabels(s.index, fontsize=8)
    ax.invert_yaxis()
    for i, v in enumerate(s.values):
        ax.text(v + 0.3, i, f"{_fmt(v)} ({_pct(v, total)}%)", va="center", fontsize=7.5, color=H_TEXT)
    ax.set_xlabel("Consultes", fontsize=8)
    fig.tight_layout(pad=0.6)
    return _buf(fig)

def _chart_canal(canals, total):
    _mpl_style()
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(11, 3.5))
    colors_list = CHART_COLORS[:len(canals)]
    wedges, texts, autotexts = ax1.pie(
        canals.values, labels=canals.index,
        colors=colors_list, autopct="%1.1f%%",
        startangle=90, pctdistance=0.75,
        textprops={"fontsize": 7.5},
    )
    for at in autotexts:
        at.set_color("white"); at.set_fontsize(7)
    ax1.set_title("Distribució per canal", fontsize=9, fontweight="bold", color=H_NAVY, pad=6)
    ax2.barh(range(len(canals)), canals.values, color=colors_list, alpha=0.85, height=0.6)
    ax2.set_yticks(range(len(canals)))
    ax2.set_yticklabels(canals.index, fontsize=8)
    ax2.invert_yaxis()
    for i, v in enumerate(canals.values):
        ax2.text(v + 0.3, i, f"{_fmt(v)}", va="center", fontsize=7.5, color=H_TEXT)
    ax2.set_xlabel("Consultes", fontsize=8)
    ax2.set_title("Volum per canal", fontsize=9, fontweight="bold", color=H_NAVY, pad=6)
    fig.tight_layout(pad=0.6)
    return _buf(fig)

def _chart_temps(df):
    col = "val_hours_resolution"
    if col not in df.columns:
        return None
    mask_enc = pd.to_numeric(
        df.get("val_encuestable", pd.Series(1, index=df.index)), errors="coerce"
    ) == 1
    serie = pd.to_numeric(df.loc[mask_enc, col], errors="coerce").dropna() * 60
    if serie.empty:
        return None
    _mpl_style()
    fig, ax = plt.subplots(figsize=(10, 3.2))
    p95 = serie.quantile(0.95)
    clipped = serie[serie <= p95]
    ax.hist(clipped, bins=40, color=H_AMBER, alpha=0.85, edgecolor="white", linewidth=0.5)
    ax.axvline(serie.median(), color=H_NAVY, linewidth=1.5, linestyle="--",
               label=f"Mediana: {_fmt_time(serie.median())}")
    ax.axvline(serie.mean(), color=H_TEAL, linewidth=1.5, linestyle="-",
               label=f"Mitjana: {_fmt_time(serie.mean())}")
    ax.set_xlabel("Minuts", fontsize=8)
    ax.set_ylabel("Freqüència", fontsize=8)
    ax.set_title("Distribució del temps de resolució (fins P95)", fontsize=9,
                 fontweight="bold", color=H_NAVY, pad=6)
    ax.legend(fontsize=7.5)
    fig.tight_layout(pad=0.6)
    return _buf(fig)

def _chart_satisfaccio(df):
    col_enc = "val_media_enc"
    cols_q   = ["val_pregunta1", "val_pregunta2", "val_pregunta3"]
    q_labels = ["P1 — Facilitat gestió", "P2 — Temps d'espera", "P3 — Tracte rebut"]
    for c in cols_q + [col_enc]:
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce")

    enc = df[col_enc].dropna() if col_enc in df.columns else pd.Series(dtype=float)
    _mpl_style()
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 3.8))
    # Distribució general
    if not enc.empty:
        bins = [0.5, 1.5, 2.5, 3.5, 4.5, 5.5]
        labels_hist = ["1", "2", "3", "4", "5"]
        counts, _ = np.histogram(enc, bins=bins)
        colors_hist = [H_RED if i < 2 else H_AMBER if i == 2 else H_TEAL for i in range(5)]
        ax1.bar(labels_hist, counts, color=colors_hist, alpha=0.85, edgecolor="white")
        ax1.axvline(enc.mean() - 1, color=H_NAVY, linewidth=1.5, linestyle="--",
                    label=f"Mitjana: {enc.mean():.2f}")
        ax1.set_xlabel("Puntuació (1-5)", fontsize=8)
        ax1.set_ylabel("Nombre enquestes", fontsize=8)
        ax1.set_title("Distribució valoració global", fontsize=9,
                      fontweight="bold", color=H_NAVY, pad=6)
        ax1.legend(fontsize=7.5)
    else:
        ax1.text(0.5, 0.5, "Sense dades", ha="center", va="center",
                 transform=ax1.transAxes, color=H_MUTED, fontsize=10)

    # Per preguntes
    q_means, q_ns = [], []
    has_q = False
    for c in cols_q:
        if c in df.columns:
            vals = df[c].dropna()
            if len(vals) > 0:
                q_means.append(vals.mean()); q_ns.append(len(vals)); has_q = True
            else:
                q_means.append(None); q_ns.append(0)
        else:
            q_means.append(None); q_ns.append(0)

    if has_q:
        valid = [(lbl, m, n) for lbl, m, n in zip(q_labels, q_means, q_ns) if m is not None]
        lbls, means, ns = zip(*valid) if valid else ([], [], [])
        y = range(len(lbls))
        ax2.barh(list(y), means, color=H_AMBER, alpha=0.85, height=0.5)
        ax2.set_xlim(0, 5.8)
        ax2.axvline(5, color=H_MID, linewidth=0.8, linestyle="--")
        for i, (m, n) in enumerate(zip(means, ns)):
            ax2.text(m + 0.08, i, f"{m:.2f}  (n={n})", va="center", fontsize=8, color=H_TEXT)
        ax2.set_yticks(list(y))
        ax2.set_yticklabels(lbls, fontsize=8)
        ax2.set_xlabel("Puntuació mitjana (1–5)", fontsize=8)
        ax2.set_title("Puntuació per Preguntes", fontsize=9,
                      fontweight="bold", color=H_NAVY, pad=6)
    else:
        ax2.text(0.5, 0.5, "Sense dades", ha="center", va="center",
                 transform=ax2.transAxes, color=H_MUTED, fontsize=10)

    fig.tight_layout(pad=0.8)
    return _buf(fig)


# ── Conclusions automàtiques ──────────────────────────────────────────────────

def _build_conclusions(doc, df, total, avg_min, enc_mean, enc_n, flg_res, dias):
    """
    Genera fins a 4 conclusions en català, ordenades per impacte.
    [ATENCIÓ] vermell · [POSITIU] verd · [INFORMATIU] blau · [SENSE DADES] gris
    """
    LABEL_COLOR = {
        "ATENCIÓ":     ("C0392B", RGBColor(0xC0, 0x39, 0x2B)),
        "POSITIU":     ("1A7A1A", RGBColor(0x1A, 0x7A, 0x1A)),
        "INFORMATIU":  ("1B3A6B", RGBColor(0x1B, 0x3A, 0x6B)),
        "SENSE DADES": ("7A8FA6", RGBColor(0x7A, 0x8F, 0xA6)),
    }
    LABEL_ORDER = {"ATENCIÓ": 0, "POSITIU": 1, "INFORMATIU": 2, "SENSE DADES": 3}
    pool = []

    # Satisfacció
    if enc_n and enc_mean is not None:
        enc_str = f"{enc_mean:.2f}".replace(".", ",")
        if enc_mean < 3.5:
            pool.append((0, "ATENCIÓ",
                f"La valoració mitjana és de {enc_str}/5 sobre {_fmt(enc_n)} enquestes, "
                f"per sota del llindar acceptable (3,5). Cal revisar la qualitat del servei "
                f"i identificar els principals punts de fricció amb el ciutadà."))
        elif enc_mean >= 4.5:
            pool.append((1, "POSITIU",
                f"Excel·lent valoració del servei: {enc_str}/5 sobre {_fmt(enc_n)} enquestes rebudes. "
                f"Indica un alt grau de satisfacció ciutadana amb l'atenció prestada."))
        elif enc_mean < 4.0:
            pool.append((2, "INFORMATIU",
                f"La valoració mitjana és de {enc_str}/5 ({_fmt(enc_n)} enquestes). "
                f"Hi ha marge de millora per assolir nivells excel·lents (≥ 4,5)."))
        else:
            pool.append((2, "INFORMATIU",
                f"Bona valoració del servei: {enc_str}/5 sobre {_fmt(enc_n)} enquestes. Nivell satisfactori."))
    else:
        pool.append((3, "SENSE DADES",
            "No s'han rebut enquestes de valoració durant aquest període. "
            "No és possible avaluar la satisfacció del servei."))

    # Temps de resolució
    if avg_min is not None:
        if avg_min > 30:
            pool.append((0, "ATENCIÓ",
                f"El temps mitjà de resolució és de {_fmt_time(avg_min)}, superior als 30 minuts recomanats. "
                f"Pot indicar acumulació de consultes complexes o necessitat de reforç en la capacitat de resposta."))
        elif avg_min < 8:
            pool.append((1, "POSITIU",
                f"Resolució molt ràpida: temps mitjà de {_fmt_time(avg_min)}. "
                f"Indica una gestió eficient i àgil de les consultes ciutadanes."))
        else:
            pool.append((2, "INFORMATIU",
                f"El temps mitjà de resolució és de {_fmt_time(avg_min)}, dins dels paràmetres habituals del servei."))
    else:
        pool.append((3, "SENSE DADES",
            "No hi ha dades de temps de resolució disponibles per calcular la durada mitjana de les consultes."))

    # Canal dominant
    if "des_entry_channel" in df.columns and total > 0:
        canals = df["des_entry_channel"].value_counts()
        canal_top = canals.index[0]
        canal_pct = _pct(canals.iloc[0], total)
        if canal_pct >= 80:
            pool.append((2, "INFORMATIU",
                f"El canal {canal_top} concentra el {canal_pct}% de les consultes. Distribució molt desequilibrada; "
                f"cal valorar si s'estan potenciant adequadament els canals alternatius."))
        elif canal_pct >= 60:
            pool.append((2, "INFORMATIU",
                f"El canal dominant és {canal_top} amb el {canal_pct}% del total, "
                f"amb presència complementària d'altres canals."))

    # Categoria dominant
    if "des_category_0" in df.columns and total > 0:
        cats = df["des_category_0"].value_counts()
        cat_top = cats.index[0]
        cat_pct = _pct(cats.iloc[0], total)
        if cat_pct >= 50:
            pool.append((2, "INFORMATIU",
                f"La categoria '{cat_top}' representa el {cat_pct}% de totes les consultes, "
                f"evidenciant una demanda molt concentrada en un únic àmbit temàtic."))

    # Variabilitat diària
    if len(dias) >= 5:
        max_d = int(dias.max())
        min_d = int(dias[dias > 0].min()) if (dias > 0).any() else 1
        ratio = max_d / min_d if min_d else 0
        if ratio >= 4:
            pool.append((2, "INFORMATIU",
                f"Alta variabilitat diària: el dia de major afluència va registrar {_fmt(max_d)} consultes, "
                f"{ratio:.0f}× més que el de menor activitat ({_fmt(min_d)}). "
                f"Pot indicar pics puntuals o dies festius."))

    # Ràtio resolució
    if flg_res is not None:
        flg_str = str(flg_res).replace(".", ",")
        if flg_res < 50:
            pool.append((0, "ATENCIÓ",
                f"El ràtio de resolució és del {flg_str}%, per sota del 50%. "
                f"Moltes consultes requereixen seguiment posterior, augmentant la càrrega operativa."))
        elif flg_res >= 80:
            pool.append((1, "POSITIU",
                f"Alt ràtio de resolució: el {flg_str}% de les consultes es tanquen satisfactòriament. "
                f"Indica una bona capacitat de resposta del servei."))

    pool.sort(key=lambda x: (LABEL_ORDER[x[1]], x[0]))
    pool = pool[:4]
    if not pool:
        return

    _section_title(doc, "7", "Conclusions i Observacions")
    _body_para(doc, "Resum automàtic de les principals observacions del període, ordenades per impacte:")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    for _, label, text in pool:
        hex_bg, rgb_color = LABEL_COLOR[label]
        # Taula d'1 fila, 2 columnes: badge | text
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        # Cel·la badge
        badge_cell = tbl.rows[0].cells[0]
        badge_cell.width = Cm(2.4)
        _set_cell_bg(badge_cell, hex_bg)
        bp = badge_cell.paragraphs[0]
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bp.paragraph_format.space_before = Pt(4)
        bp.paragraph_format.space_after  = Pt(4)
        br = bp.add_run(label)
        br.font.bold  = True
        br.font.size  = Pt(8)
        br.font.color.rgb = C_WHITE
        br.font.name  = "Calibri"
        badge_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Cel·la text
        text_cell = tbl.rows[0].cells[1]
        text_cell.width = Cm(14.6)
        _set_cell_bg(text_cell, "FFFFFF")
        _set_cell_borders(text_cell, left={"val": "single", "sz": 6, "color": hex_bg})
        tp = text_cell.paragraphs[0]
        tp.paragraph_format.space_before = Pt(4)
        tp.paragraph_format.space_after  = Pt(4)
        tr = tp.add_run(text)
        tr.font.size  = Pt(9.5)
        tr.font.color.rgb = C_TEXT
        tr.font.name  = "Calibri"
        text_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        doc.add_paragraph().paragraph_format.space_after = Pt(3)


# ── Portada ───────────────────────────────────────────────────────────────────

def _add_cover(doc, periode_label, generated):
    """Pàgina de portada."""
    # Espai superior
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Títol principal
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(10)
    r = p_title.add_run("Informe Mensual de Consultes")
    r.font.bold  = True
    r.font.size  = Pt(30)
    r.font.color.rgb = C_NAVY
    r.font.name  = "Calibri"

    # Subtítol període
    p_per = doc.add_paragraph()
    p_per.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_per.paragraph_format.space_after = Pt(8)
    r2 = p_per.add_run(periode_label)
    r2.font.size  = Pt(18)
    r2.font.color.rgb = C_TEAL
    r2.font.name  = "Calibri"

    # Servei
    p_serv = doc.add_paragraph()
    p_serv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_serv.paragraph_format.space_after = Pt(4)
    r3 = p_serv.add_run("OAC 360º · Adtende Analytics")
    r3.font.size  = Pt(12)
    r3.font.color.rgb = C_MUTED
    r3.font.name  = "Calibri"

    _add_hrule(doc, color=_hex(C_TEAL), size=8)

    for _ in range(10):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Data generació
    p_gen = doc.add_paragraph()
    p_gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p_gen.add_run(f"Generat el {generated}")
    r4.font.size  = Pt(9)
    r4.font.color.rgb = C_MUTED
    r4.font.name  = "Calibri"

    doc.add_page_break()


# ── PDF builder → DOCX builder ────────────────────────────────────────────────

def _build_docx(df, year, month, periode_label, output_path):
    # Forcem tipus numèrics
    for col in ["val_rating", "val_pregunta1", "val_pregunta2", "val_pregunta3",
                "val_encuestable", "val_time_spent", "val_hours_resolution",
                "val_hours_spent", "flg_resolution", "val_population",
                "val_media_enc", "val_threads", "flg_service_hour",
                "val_time_spent_last_working_day"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")

    generated = date.today().strftime("%d/%m/%Y")
    doc = Document()

    # Marges de pàgina
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── PORTADA ───────────────────────────────────────────────────────────────
    _add_cover(doc, periode_label, generated)

    # ── Mètriques ─────────────────────────────────────────────────────────────
    total = len(df)

    avg_rating, n_rat = None, 0
    if "val_rating" in df.columns:
        r = pd.to_numeric(df["val_rating"], errors="coerce").dropna()
        n_rat = len(r)
        if n_rat: avg_rating = round(r.mean(), 2)

    avg_min = None
    if "val_hours_resolution" in df.columns:
        mask_enc = pd.to_numeric(
            df.get("val_encuestable", pd.Series(1, index=df.index)), errors="coerce"
        ) == 1
        t = pd.to_numeric(df.loc[mask_enc, "val_hours_resolution"], errors="coerce").dropna()
        if len(t): avg_min = t.mean() * 60

    flg_res = None
    if "flg_resolution" in df.columns:
        res = pd.to_numeric(df["flg_resolution"], errors="coerce")
        flg_res = round(res.sum() / len(res) * 100, 1) if len(res) else None

    n_cats   = df["des_category_0"].nunique() if "des_category_0" in df.columns else 0
    n_canals = df["des_entry_channel"].nunique() if "des_entry_channel" in df.columns else 0

    import calendar as _cal
    _date_col = "td_managed" if "td_managed" in df.columns else "td_created"
    dias = df.groupby(pd.to_datetime(df[_date_col], errors="coerce").dt.date).size()
    dies_mes = _cal.monthrange(year, month)[1]
    max_dia = int(dias.max()) if len(dias) else 0
    min_dia = int(dias.min()) if len(dias) else 0

    col_enc = "val_media_enc"
    enc_series = pd.to_numeric(df.get(col_enc, pd.Series(dtype=float)), errors="coerce").dropna()
    n_enc = len(enc_series)
    enc_mean = round(enc_series.mean(), 2) if n_enc > 0 else None

    # ── KPIs ─────────────────────────────────────────────────────────────────
    kpis = [(_fmt(total), "Total Consultes", "#1B3A6B")]
    if flg_res is not None:
        kpis.append((f"{str(flg_res).replace('.',',')}%", "Ratio Resolució", "#00B4A6"))
    if avg_min is not None:
        kpis.append((_fmt_time(avg_min), "Temps Resolució", "#F4A620"))
    if enc_mean is not None:
        kpis.append((f"{enc_mean:.2f}".replace(".",","), "Valoració Mitjana", "#00B4A6"))
    kpis.append((_fmt(n_cats),   "Categories", "#1B3A6B"))
    kpis.append((_fmt(n_canals), "Canals",     "#1B3A6B"))

    _add_kpi_row(doc, kpis)

    # ── SECCIÓ 1: Evolució diària ─────────────────────────────────────────────
    _section_title(doc, "1", "Evolució Diària de Consultes")
    _body_para(doc,
        f"Al llarg del mes es van registrar {_fmt(total)} consultes en {dies_mes} dies. "
        f"El dia de major activitat va tenir {_fmt(max_dia)} consultes "
        f"i el de menor activitat {_fmt(min_dia)}.")
    chart = _chart_evolucion(df)
    if chart:
        doc.add_picture(chart, width=Cm(16))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── SECCIÓ 2: Distribució dia setmana i hora ──────────────────────────────
    _section_title(doc, "2", "Distribució per Dia Setmana i Hora")
    _body_para(doc,
        "El mapa de calor mostra la concentració de consultes per cada combinació "
        "de dia de la setmana i hora del dia. Les cel·les més fosques indiquen major activitat.")
    chart = _chart_heatmap(df)
    if chart:
        doc.add_picture(chart, width=Cm(16))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if "td_created" in df.columns:
        df2 = df.copy()
        df2["td_created"] = pd.to_datetime(df2["td_created"], errors="coerce")
        df2["dia_num"] = df2["td_created"].dt.dayofweek
        per_dia = df2.groupby("dia_num").size().reindex(range(7), fill_value=0)
        rows_dia = [[DIES[i], _fmt(v), f"{_pct(v, total)}%"] for i, v in per_dia.items()]
        rows_dia.append(["TOTAL", _fmt(total), "100%"])
        _add_table(doc, ["Dia", "Consultes", "% Total"], rows_dia,
                   col_widths_cm=[7, 4, 4], total_row=True)

    # ── SECCIÓ 3: Distribució per categoria ──────────────────────────────────
    _section_title(doc, "3", "Distribució per Categoria de Consultes")
    if "des_category_0" in df.columns:
        cats = df["des_category_0"].value_counts()
        top  = cats.index[0]
        _body_para(doc,
            f"S'han identificat {n_cats} categories. "
            f"La més freqüent és '{top}' amb {_fmt(cats.iloc[0])} consultes "
            f"({_pct(cats.iloc[0], total)}% del total).")
        chart = _chart_hbars(cats, total, color=H_NAVY)
        if chart:
            doc.add_picture(chart, width=Cm(16))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        rows_cat = [[cat, _fmt(cnt), f"{_pct(cnt, total)}%"] for cat, cnt in cats.items()]
        rows_cat.append(["TOTAL", _fmt(total), "100%"])
        _add_table(doc, ["Categoria", "Consultes", "% Total"], rows_cat,
                   col_widths_cm=[10, 3.5, 3.5], total_row=True)

    # ── SECCIÓ 4: Canal d'assistència ────────────────────────────────────────
    _section_title(doc, "4", "Distribució per Canal d'Assistència")
    if "des_entry_channel" in df.columns:
        canals = df["des_entry_channel"].value_counts()
        _body_para(doc,
            f"Les consultes s'han atès a través de {n_canals} canals. "
            f"El canal principal és '{canals.index[0]}' "
            f"amb el {_pct(canals.iloc[0], total)}% del total.")
        chart = _chart_canal(canals, total)
        if chart:
            doc.add_picture(chart, width=Cm(16))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        rows_can = [[c, _fmt(v), f"{_pct(v, total)}%"] for c, v in canals.items()]
        rows_can.append(["TOTAL", _fmt(total), "100%"])
        _add_table(doc, ["Canal", "Consultes", "% Total"], rows_can,
                   col_widths_cm=[10, 3.5, 3.5], total_row=True)

    # ── SECCIÓ 5: Temps de resolució ─────────────────────────────────────────
    _section_title(doc, "5", "Temps de Resolució")
    if avg_min is not None:
        mask_enc = pd.to_numeric(
            df.get("val_encuestable", pd.Series(1, index=df.index)), errors="coerce"
        ) == 1
        serie_min = pd.to_numeric(
            df.loc[mask_enc, "val_hours_resolution"], errors="coerce"
        ).dropna() * 60
        p25 = serie_min.quantile(0.25)
        p50 = serie_min.quantile(0.50)
        p75 = serie_min.quantile(0.75)
        p95 = serie_min.quantile(0.95)
        _body_para(doc,
            f"La mitjana de temps de resolució és de {_fmt_time(avg_min)} "
            f"(calculat sobre {_fmt(len(serie_min))} consultes encuestables). "
            f"La meitat es resolen en menys de {_fmt_time(p50)} (mediana).")
        chart = _chart_temps(df)
        if chart:
            doc.add_picture(chart, width=Cm(15))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        rows_t = [
            ["Percentil 25",  _fmt_time(p25)],
            ["Mediana (P50)", _fmt_time(p50)],
            ["Mitjana",       _fmt_time(avg_min)],
            ["Percentil 75",  _fmt_time(p75)],
            ["Percentil 95",  _fmt_time(p95)],
        ]
        _add_table(doc, ["Estadístic", "Valor"], rows_t, col_widths_cm=[9, 8])
    else:
        _body_para(doc, "No hi ha dades de temps disponibles.")

    # ── SECCIÓ 6: Satisfacció i valoració ────────────────────────────────────
    _section_title(doc, "6", "Satisfacció i Valoració")
    if n_enc > 0:
        enc_median = round(enc_series.median(), 2)
        n_encuestable = int(pd.to_numeric(
            df.get("val_encuestable", pd.Series(dtype=float)), errors="coerce"
        ).sum()) if "val_encuestable" in df.columns else total
        taxa = round(n_enc / n_encuestable * 100, 1) if n_encuestable else None

        enc_kpis = [
            (f"{enc_mean:.2f}".replace(".", ","),    "Valoració Mitjana",  "#00B4A6"),
            (f"{enc_median:.2f}".replace(".", ","),  "Valoració Mediana",  "#1B3A6B"),
            (_fmt(n_enc),                            "Enquestes Rebudes",  "#F4A620"),
        ]
        if taxa is not None:
            enc_kpis.append((f"{str(taxa).replace('.', ',')}%", "Taxa Resposta", "#1B3A6B"))
        _add_kpi_row(doc, enc_kpis)

        chart_sat = _chart_satisfaccio(df)
        if chart_sat:
            doc.add_picture(chart_sat, width=Cm(16))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        if "des_entry_channel" in df.columns:
            sat_canal = (
                df[df[col_enc].notna()]
                .groupby("des_entry_channel")[col_enc]
                .agg(["mean", "count"])
                .rename(columns={"mean": "Mitjana", "count": "Enquestes"})
                .query("Enquestes > 0")
                .sort_values("Enquestes", ascending=False)
            )
            if not sat_canal.empty:
                _body_para(doc, "Satisfacció per canal d'assistència:")
                rows_sat = [
                    [canal, f"{row['Mitjana']:.2f}".replace(".", ","), _fmt(int(row["Enquestes"]))]
                    for canal, row in sat_canal.iterrows()
                ]
                _add_table(doc, ["Canal", "Valoració Mitjana", "Enquestes"], rows_sat,
                           col_widths_cm=[9, 4.5, 3.5])
    else:
        _body_para(doc, "No hi ha dades de valoració disponibles per aquest període.")

    # ── SECCIÓ 7: Conclusions automàtiques ───────────────────────────────────
    _build_conclusions(doc, df, total, avg_min, enc_mean, n_enc, flg_res, dias)

    doc.save(str(output_path))


# ── Filtre doble de data (igual que el BI) ────────────────────────────────────

def _apply_dual_date_filter(df, year, month):
    """
    El BI aplica SEMPRE dos filtres simultanis: td_managed I td_created
    han d'estar AMBDÓS dins del rang del mes.
    Comparació com strings YYYY-MM-DD per evitar pèrdua de registres amb
    dates no parseables (pd.to_datetime genera NaT silenciosament).
    """
    if month == 12:
        d_to = f"{year + 1}-01-01"
    else:
        d_to = f"{year}-{month + 1:02d}-01"
    d_from = f"{year}-{month:02d}-01"

    mask = pd.Series(True, index=df.index)
    if "td_managed" in df.columns:
        mask &= (df["td_managed"] >= d_from) & (df["td_managed"] < d_to)
    if "td_created" in df.columns:
        mask &= (df["td_created"] >= d_from) & (df["td_created"] < d_to)
    return df[mask].reset_index(drop=True)


# ── Configuració dels 5 serveis ───────────────────────────────────────────────

SERVICES = [
    {
        "key":        "oac360",
        "name":       "OAC 360",
        "endpoint":   "tickets_enriquits",
        "project":    "OAC 360º",
        "date_field": "td_managed",
        "slug":       "OAC360",
    },
    {
        "key":        "oac360_social",
        "name":       "OAC 360 Social",
        "endpoint":   "tickets_enriquits",
        "project":    "OAC 360º SOCIAL",
        "date_field": "td_managed",
        "slug":       "OAC360_Social",
    },
    {
        "key":        "oac360_tributs",
        "name":       "OAC 360 Tributs",
        "endpoint":   "tickets_enriquits",
        "project":    "OAC 360º Tributs",
        "date_field": "td_managed",
        "slug":       "OAC360_Tributs",
    },
    {
        "key":        "satediba",
        "name":       "SATE DIBA",
        "endpoint":   "tickets_enriquits",
        "project":    "SATE DIBA",
        "date_field": "td_managed",
        "slug":       "SATEDIBA",
    },
    {
        "key":        "centraleta",
        "name":       "Centraleta",
        "endpoint":   "tickets_enriquits",
        "project":    "Centraleta",
        "date_field": "td_managed",
        "slug":       "Centraleta",
    },
]


# ── Funcions principals ───────────────────────────────────────────────────────

def generate_monthly_report(year: int, month: int,
                             endpoint: str = "tickets_enriquits",
                             date_field: str = "td_managed",
                             project: str = "OAC 360º",
                             output_dir: str = "."):
    client = AdtendeClient()
    print("Autenticant...")
    client.login()

    mes = MESOS[month]
    print(f"Descarregant {mes} {year} (filtre: {date_field}, projecte: {project})...")
    df = client.query_month(endpoint, year, month, date_field=date_field, project=project)
    print(f"  {len(df)} registres API.")
    df = _apply_dual_date_filter(df, year, month)
    print(f"  {len(df)} registres després del filtre doble (td_managed + td_created).")

    if df.empty:
        print("No hi ha dades per al període seleccionat.")
        return None

    output_path = Path(output_dir) / f"informe_{year}_{month:02d}_{mes}.docx"
    print("Generant Word...")
    _build_docx(df, year, month, f"{mes} {year}", output_path)
    print(f"Informe generat: {output_path}")
    return str(output_path)


def generate_all_monthly_reports(year: int, month: int, output_dir: str = "."):
    """
    Genera els 5 informes de servei (Word) per al mes indicat:
      - OAC 360
      - OAC 360 Social
      - OAC 360 Tributs
      - SATE DIBA
      - Centraleta
    Retorna la llista de paths generats.
    """
    client = AdtendeClient()
    print("Autenticant...")
    client.login()

    mes = MESOS[month]
    periode = f"{mes} {year}"
    out_dir = Path(output_dir)
    generated = []

    for svc in SERVICES:
        print(f"\n{'─'*50}")
        print(f"▶ {svc['name']} ({svc['endpoint']}, projecte: {svc['project'] or 'tots'})")
        try:
            df = client.query_month(
                svc["endpoint"], year, month,
                date_field=svc["date_field"],
                project=svc["project"],
            )
            print(f"  {len(df)} registres API.")
            df = _apply_dual_date_filter(df, year, month)
            print(f"  {len(df)} registres després del filtre doble (td_managed + td_created).")
            if df.empty:
                print(f"  ⚠️  Sense dades — s'omet {svc['name']}.")
                continue

            fname = f"informe_{year}_{month:02d}_{mes}_{svc['slug']}.docx"
            output_path = out_dir / fname
            title = f"{svc['name']} — {periode}"
            print(f"  Generant Word → {fname}...")
            _build_docx(df, year, month, title, output_path)
            print(f"  ✅ Generat: {output_path}")
            generated.append(str(output_path))
        except Exception as e:
            print(f"  ❌ Error en {svc['name']}: {e}")

    print(f"\n{'─'*50}")
    print(f"Total informes generats: {len(generated)}/{len(SERVICES)}")
    for p in generated:
        print(f"  · {p}")
    return generated


def generate_daily_report(year: int, month: int, day: int,
                           endpoint: str = "tickets_enriquits",
                           date_field: str = "td_managed",
                           project: str = "OAC 360º",
                           output_dir: str = "."):
    client = AdtendeClient()
    print("Autenticant...")
    client.login()

    mes = MESOS[month]
    print(f"Descarregant {day:02d} {mes} {year} (filtre: {date_field}, projecte: {project})...")
    df = client.query_date(endpoint, year, month, day, date_field=date_field, project=project)
    print(f"  {len(df)} registres descarregats.")

    if df.empty:
        print("No hi ha dades per al període seleccionat.")
        return None

    output_path = Path(output_dir) / f"informe_{year}_{month:02d}_{day:02d}.docx"
    print("Generant Word...")
    _build_docx(df, year, month, f"{day:02d} {mes} {year}", output_path)
    print(f"Informe generat: {output_path}")
    return str(output_path)


# ── Informe comparatiu ────────────────────────────────────────────────────────

def _metrics(df):
    """Extreu KPIs d'un DataFrame."""
    total = len(df)
    mask_enc = pd.to_numeric(df.get("val_encuestable", pd.Series(1, index=df.index)),
                             errors="coerce") == 1
    avg_min = None
    if "val_hours_resolution" in df.columns:
        t = pd.to_numeric(df.loc[mask_enc, "val_hours_resolution"], errors="coerce").dropna()
        if len(t): avg_min = t.mean() * 60

    avg_rating = None
    col_enc = "val_media_enc"
    if col_enc in df.columns:
        enc = pd.to_numeric(df[col_enc], errors="coerce").dropna()
        if len(enc): avg_rating = round(enc.mean(), 2)

    n_canals = df["des_entry_channel"].nunique() if "des_entry_channel" in df.columns else 0

    return {
        "total": total,
        "avg_min": avg_min,
        "avg_rating": avg_rating,
        "n_canals": n_canals,
    }


def generate_yoy_report(year1: int, month: int, year2: int,
                        endpoint: str = "tickets_enriquits",
                        date_field: str = "td_managed",
                        project: str = "OAC 360º",
                        output_dir: str = "."):
    """Informe comparatiu any sobre any (year1 vs year2) per al mes indicat."""
    client = AdtendeClient()
    print("Autenticant...")
    client.login()

    mes = MESOS[month]
    print(f"Descarregant {mes} {year1}...")
    df1 = client.query_month(endpoint, year1, month, date_field=date_field, project=project)
    print(f"  {len(df1)} registres.")
    print(f"Descarregant {mes} {year2}...")
    df2 = client.query_month(endpoint, year2, month, date_field=date_field, project=project)
    print(f"  {len(df2)} registres.")

    m1 = _metrics(df1)
    m2 = _metrics(df2)

    def _delta(v1, v2):
        if v1 is None or v2 is None or v1 == 0:
            return "—"
        d = (v2 - v1) / v1 * 100
        sign = "+" if d >= 0 else ""
        return f"{sign}{d:.1f}%"

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    generated = date.today().strftime("%d/%m/%Y")
    _add_cover(doc, f"Comparativa {mes} {year1} vs {year2}", generated)

    _section_title(doc, "1", f"Resum comparatiu — {mes}")
    rows_cmp = [
        ["Total consultes",    _fmt(m1["total"]),
         _fmt(m2["total"]),    _delta(m1["total"], m2["total"])],
        ["Temps mig resolució", _fmt_time(m1["avg_min"]),
         _fmt_time(m2["avg_min"]), "—"],
        ["Valoració mitjana",  f"{m1['avg_rating']:.2f}".replace(".", ",") if m1["avg_rating"] else "—",
         f"{m2['avg_rating']:.2f}".replace(".", ",") if m2["avg_rating"] else "—",
         _delta(m1["avg_rating"], m2["avg_rating"])],
        ["Canals actius",      str(m1["n_canals"]),
         str(m2["n_canals"]),  "—"],
    ]
    _add_table(doc,
               ["Indicador", str(year1), str(year2), "Variació"],
               rows_cmp,
               col_widths_cm=[7, 3, 3, 4])

    output_path = Path(output_dir) / f"informe_yoy_{year1}_{year2}_{month:02d}_{mes}.docx"
    doc.save(str(output_path))
    print(f"Informe comparatiu generat: {output_path}")
    return str(output_path)
